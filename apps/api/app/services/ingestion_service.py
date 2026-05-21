"""
Document ingestion service: parses, chunks, embeds, and stores documents.
Supports PDF, DOCX, TXT, Markdown, HTML, CSV with optional OCR.
Called by Inngest background function for async processing.
"""
import io
import logging
import os
import re
import tempfile
from pathlib import Path
from typing import Optional
from uuid import UUID

import httpx
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.models import Document, DocumentChunk
from app.services.embedding_service import embed_texts

logger = logging.getLogger(__name__)

# Chunk configuration
CHUNK_SIZE = 1000  # characters
CHUNK_OVERLAP = 200


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------

async def download_file(url: str) -> tuple[bytes, str]:
    """Download file from Vercel Blob URL, return (content_bytes, filename)."""
    async with httpx.AsyncClient(timeout=60) as client:
        response = await client.get(url)
        response.raise_for_status()
        content_type = response.headers.get("content-type", "")
        filename = url.split("/")[-1].split("?")[0]
        return response.content, filename


def extract_text_from_pdf(content: bytes) -> str:
    """Extract text from PDF using pdfplumber, fallback to OCR."""
    import pdfplumber

    text_parts = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text and page_text.strip():
                text_parts.append(page_text)
            else:
                # Fallback to OCR for image-based pages
                ocr_text = _ocr_page(page)
                if ocr_text:
                    text_parts.append(ocr_text)

    return "\n\n".join(text_parts)


def _ocr_page(page) -> str:
    """OCR a single PDF page using pytesseract."""
    try:
        import pytesseract
        from PIL import Image

        img = page.to_image(resolution=200).original
        return pytesseract.image_to_string(img, lang="eng")
    except Exception as e:
        logger.warning("OCR failed for page: %s", e)
        return ""


def extract_text_from_docx(content: bytes) -> str:
    """Extract text from DOCX."""
    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(content))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)
    return "\n\n".join(paragraphs)


def extract_text_from_html(content: bytes) -> str:
    """Extract text from HTML, removing scripts and styles."""
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(content, "lxml")
    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()
    text = soup.get_text(separator="\n")
    # Collapse whitespace
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    return "\n".join(lines)


def extract_text_from_csv(content: bytes) -> str:
    """Convert CSV to text representation."""
    import pandas as pd

    try:
        df = pd.read_csv(io.BytesIO(content), encoding="utf-8")
    except UnicodeDecodeError:
        df = pd.read_csv(io.BytesIO(content), encoding="latin-1")

    lines = []
    # Header description
    lines.append(f"CSV with {len(df)} rows and columns: {', '.join(df.columns)}")
    lines.append("")
    # Convert rows to text (limit to 1000 rows to avoid huge embeddings)
    for _, row in df.head(1000).iterrows():
        lines.append(" | ".join(f"{col}: {val}" for col, val in row.items() if str(val) != "nan"))
    return "\n".join(lines)


def extract_text_from_markdown(content: bytes) -> str:
    """Extract text from Markdown."""
    text = content.decode("utf-8", errors="replace")
    # Strip markdown syntax for cleaner embedding
    text = re.sub(r"#+\s", "", text)  # headings
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)  # bold
    text = re.sub(r"\*(.+?)\*", r"\1", text)  # italic
    text = re.sub(r"`(.+?)`", r"\1", text)  # inline code
    text = re.sub(r"```[\s\S]+?```", "", text)  # code blocks (remove)
    text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)  # links
    return text


def extract_text(content: bytes, file_type: str) -> str:
    """Route extraction to the right parser based on file type."""
    file_type = file_type.lower().strip(".")
    if file_type == "pdf":
        return extract_text_from_pdf(content)
    elif file_type in ("docx", "doc"):
        return extract_text_from_docx(content)
    elif file_type in ("html", "htm"):
        return extract_text_from_html(content)
    elif file_type == "csv":
        return extract_text_from_csv(content)
    elif file_type in ("md", "markdown"):
        return extract_text_from_markdown(content)
    elif file_type == "txt":
        return content.decode("utf-8", errors="replace")
    else:
        # Attempt UTF-8 decode as fallback
        return content.decode("utf-8", errors="replace")


# ---------------------------------------------------------------------------
# Chunking
# ---------------------------------------------------------------------------

def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    """
    Split text into overlapping chunks with sentence boundary awareness.
    Uses LangChain's RecursiveCharacterTextSplitter logic.
    """
    from langchain.text_splitter import RecursiveCharacterTextSplitter

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=overlap,
        separators=["\n\n", "\n", ". ", "! ", "? ", " ", ""],
        length_function=len,
    )
    return splitter.split_text(text)


# ---------------------------------------------------------------------------
# Main ingestion pipeline
# ---------------------------------------------------------------------------

async def ingest_document(
    db: AsyncSession,
    document_id: str,
) -> dict:
    """
    Full ingestion pipeline: Download → Extract → Chunk → Embed → Store.
    Called by Inngest background function.
    """
    # Load document record
    result = await db.execute(select(Document).where(Document.id == document_id))
    doc = result.scalar_one_or_none()
    if not doc:
        raise ValueError(f"Document {document_id} not found")

    # Mark as processing
    doc.status = "processing"
    await db.commit()

    try:
        # 1. Download
        if not doc.file_url:
            raise ValueError("Document has no file_url")
        content_bytes, _ = await download_file(doc.file_url)
        logger.info("Downloaded document %s (%d bytes)", document_id, len(content_bytes))

        # 2. Extract text
        text = extract_text(content_bytes, doc.file_type)
        if not text.strip():
            raise ValueError("No text could be extracted from document")
        logger.info("Extracted %d chars from document %s", len(text), document_id)

        # 3. Chunk
        chunks = chunk_text(text)
        logger.info("Created %d chunks for document %s", len(chunks), document_id)

        # 4. Embed (batch in groups of 100 to respect rate limits)
        all_embeddings: list[list[float]] = []
        batch_size = 100
        for i in range(0, len(chunks), batch_size):
            batch = chunks[i : i + batch_size]
            batch_embeddings = await embed_texts(batch)
            all_embeddings.extend(batch_embeddings)

        # 5. Store chunks
        chunk_objects = []
        for i, (chunk_text_val, embedding) in enumerate(zip(chunks, all_embeddings)):
            chunk = DocumentChunk(
                document_id=doc.id,
                org_id=doc.org_id,
                content=chunk_text_val,
                chunk_index=i,
                embedding=embedding,
                chunk_metadata={
                    "char_count": len(chunk_text_val),
                    "chunk_index": i,
                    "total_chunks": len(chunks),
                },
            )
            chunk_objects.append(chunk)

        db.add_all(chunk_objects)

        # 6. Update document status
        doc.status = "completed"
        doc.chunk_count = len(chunks)
        await db.commit()

        logger.info(
            "Ingestion complete for document %s: %d chunks",
            document_id,
            len(chunks),
        )
        return {"document_id": document_id, "chunks": len(chunks), "status": "completed"}

    except Exception as exc:
        doc.status = "failed"
        doc.error_message = str(exc)[:500]
        await db.commit()
        logger.error("Ingestion failed for document %s: %s", document_id, exc)
        raise
